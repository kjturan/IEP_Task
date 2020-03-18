import sys
import subprocess
import re
import docx
from datetime import datetime
from dateutil.relativedelta import *
from tkinter import *
from flask import Flask
from flask_restful import Api, Resource
from fhir_parser.fhir import FHIR 

def write_all_UUID():
    file_path = "UUID.txt"
    id_file = codecs.open(file_path, "w", "utf-8-sig")
    patients = fhir.get_all_patients()
    id_list = list()
    for patient in patients:
        id_list.append(patient.uuid)
    id_string = '\n'.join(id_list)
    id_file.write(id_string)
    id_file.flush()

def get_comms(patient):
    comms = patient.communications
    langs = comms.languages
    codes = comms.codes
    combined = []
    for i in range(len(langs)):
        temp_str = langs[i] + ' [' + codes[i] + ']'
        combined.append(temp_str)
    return ', '.join(combined)

def get_addresses(patient):
    adds = patient.addresses
    adds = [str(i) for i in adds]
    formatter = '\n' + ' '*(len('Address: ')+7 )
    temp = adds[0].split('\n') 
    temp = [i + formatter for i in temp]
    return temp

def get_telecoms(patient):
    teles = patient.telecoms
    combined = []
    for tele in teles:
        use = tele.use or ''
        if use == 'home':
            use = 'Home'
        if use == 'cell':
            use = 'Cell'
        sys = tele.system or ''
        num = tele.number or ''
        temp_str = use + sys + ': ' + num
        combined.append(temp_str)
    formatter = (' '*(len('Telecoms: ')+25))        
    return ', '.join(combined)

def get_identifiers(patient):
    identifiers = patient.identifiers
    identifiers = [(str(i)) for i in identifiers if 'Medical' not in str(i)]
    identifiers = [(shorten(i)) for i in identifiers]
    return identifiers

def shorten(string):
    string = string.replace("Driver\'s License",'DL:') 
    string = string.replace('Passport Number','PPN:')
    string = string.replace('Social Security Number','SSN:')
    string = string.replace('Medical Record Number','MRN:')
    return string

def get_mutliple_birth(patient):
    try:
        boolean = patient.mutliple_birth
        if False:
            return 'False'
        else:
            return 'True'
    except:
        return 'UNK'

def calc_age(patient):
    patient_DT = patient.birth_date
    current_DT = datetime.today()
    patient_age = relativedelta(current_DT, patient_DT).years
    return str(patient_age)

def format_datetime(dt):
    return dt.strftime('%d/%m/%y')

def get_marital(patient):
    marital = patient.marital_status
    return str(marital)

def get_daly(patient):
    num = patient.get_extension('disability-adjusted-life-years') 
    return str(round(num,2))

def get_gender(patient):
    gender = patient.gender
    if gender == 'female':
        return 'F'
    if gender == 'male':
        return 'M'  
    return 'UNK'

def format_observation(observation):
    dictionary = dict()
    dictionary['UUID'] = observation.uuid
    dictionary['Encounter ID'] = observation.encounter_uuid
    dictionary['Type'] = observation.type
    dictionary['Status'] = observation.status
    dictionary['Date'] = format_datetime(observation.effective_datetime)
    return dictionary

def format_components(observation):
    components = observation.components
    comps = []
    for comp in components:
        temp = []
        temp.append(comp.system + ' (code: ' + comp.code + ')')
        temp.append(comp.display)
        temp.append(comp.quantity())
        comps.append(temp)
    return comps

def doc_name_gen(patient):
    f_name = patient.name.given
    initial = f_name[0]
    l_name = patient.name.family
    doc_str = initial + l_name
    return doc_str

def convert_to(folder, source, timeout=None):
    args = [libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', folder, source]
    process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
    filename = re.search('-> (.*?) using filter', process.stdout.decode())
    return filename.group(1)

def libreoffice_exec():
    if sys.platform == 'darwin':
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    return 'libreoffice'

def write_word(patient,observations):
    document = docx.Document()
    header = document.add_heading('Patient Info', 0)

    header = document.add_heading('Identifiers', 1)
    identifiers = get_identifiers(patient)
    p = document.add_paragraph('UUID: ' + identifiers[0])
    p = document.add_paragraph(identifiers[1])
    p = document.add_paragraph(identifiers[2])
    p = document.add_paragraph(identifiers[3])

    header = document.add_heading('Personal Information', 1).underline = True
    table = document.add_table(rows=5, cols=2)
    table.allow_autofit = True

    cols = table.columns[0]
    p = cols.cells[0].add_paragraph('Name: ' + patient.name.given + patient.name.family)
    p = cols.cells[1].add_paragraph('Title: ' + patient.name.prefix)
    p = cols.cells[2].add_paragraph('Birthdate: ' + format_datetime(patient.birth_date))
    p.add_run(' (age: ' + calc_age(patient) + ')')
    p = cols.cells[3].add_paragraph('Race: ' + patient.get_extension('us-core-race'))
    p = cols.cells[4].add_paragraph('Ethnicity: ' + patient.get_extension('us-core-ethnicity'))

    cols = table.columns[1]
    p = cols.cells[0].add_paragraph('Gender: ' + get_gender(patient))
    p.add_run(' (Birthsex: ' + patient.get_extension('us-core-birthsex') + ')')
    p = cols.cells[1].add_paragraph('Marital Status: ')
    p.add_run(get_marital(patient))
    p = cols.cells[2].add_paragraph('Multiple Birth: ')
    p.add_run(get_mutliple_birth(patient))


    header = document.add_heading('Communication Details', 1).underline = True
    p = document.add_paragraph(get_telecoms(patient))
    p = document.add_paragraph('Communications: ')
    p.add_run(get_comms(patient))
    formatted_address = get_addresses(patient)
    p = document.add_paragraph('Address: ' + formatted_address[0]) 
    for i in range(len(formatted_address)):
        if i == 0:
            continue
        p.add_run(formatted_address[i])
    
    document.add_page_break()

    header = document.add_heading('Observation History', 0)
    for obs in observations:
        obs_dict = format_observation(obs)
        obs_keys = obs_dict.keys()
        comp_list = format_components(obs)
        for key in obs_keys:
            p = document.add_paragraph(key + ': ' + obs_dict[key])
        for comp in comp_list:
            p = document.add_paragraph('System: ' + comp[0])
            p = document.add_paragraph('Observation: ' + comp[1]).bold = True
            p = document.add_paragraph('Meassured: ' + comp[2]).bold = True
            p = document.add_paragraph()
        p = document.add_paragraph()
    doc_name = doc_name_gen(patient) + '.docx'
    document.save(doc_name)
    return doc_name